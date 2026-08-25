"""Build the SISCA UAT technical guide as a Word document for partner delivery."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "guia-sisca-preparacion-uat.md"
OUTPUT = ROOT / "deliverables" / "sisca-uat" / "Guia-preparacion-UAT-SISCA-Carobra-Rewards.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = "5B6573"


def _set_run_font(run, *, name: str = "Calibri", size: float | None = None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def _set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, *, top=60, start=120, bottom=60, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[float]):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_pr = table._tbl.tblPr
    tbl_layout = table_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        table_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_w = table_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = table_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        table_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index])
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(round(widths[index] * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_borders(table):
    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "D9E2F3")


def _keep_table_rows_intact(table):
    for index, row in enumerate(table.rows):
        row_pr = row._tr.get_or_add_trPr()
        cannot_split = OxmlElement("w:cantSplit")
        row_pr.append(cannot_split)
        if index == 0:
            repeat_header = OxmlElement("w:tblHeader")
            repeat_header.set(qn("w:val"), "true")
            row_pr.append(repeat_header)


def _append_page_number(paragraph):
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def _configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Code" not in [style.name for style in doc.styles]:
        code = doc.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = doc.styles["Code"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(9)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(4)
    code.paragraph_format.left_indent = Inches(0.18)


def _configure_section(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run("Carobra Rewards  |  Preparación UAT SISCA")
    _set_run_font(header_run, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Uso UAT — Página ")
    _set_run_font(footer_run, size=9, color=MUTED)
    _append_page_number(footer)


def _add_masthead(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("GUÍA DE PREPARACIÓN")
    _set_run_font(run, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("UAT SISCA ↔ Carobra Rewards")
    _set_run_font(run, size=23, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("Acciones requeridas de SISCA para habilitar la conexión")
    _set_run_font(run, size=13, color=MUTED)

    meta = doc.add_table(rows=2, cols=2)
    _set_table_geometry(meta, [1.35, 5.15])
    _set_borders(meta)
    entries = (
        ("Destinatario", "Equipo de integración SISCA"),
        ("Objetivo", "Habilitar UAT antes de la prueba de 100 clientes sintéticos"),
    )
    for row, (label, value) in zip(meta.rows, entries, strict=True):
        _set_cell_shading(row.cells[0], LIGHT_BLUE)
        for cell in row.cells:
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
        label_run = row.cells[0].paragraphs[0].add_run(label)
        _set_run_font(label_run, size=10, color=INK, bold=True)
        value_run = row.cells[1].paragraphs[0].add_run(value)
        _set_run_font(value_run, size=10, color="000000")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _add_table(doc: Document, lines: list[str]):
    rows = [[part.strip() for part in line.strip().strip("|").split("|")] for line in lines]
    rows.pop(1)
    columns = len(rows[0])
    table = doc.add_table(rows=0, cols=columns)
    widths = [6.5 / columns] * columns
    if columns == 2:
        widths = [2.05, 4.45]
    elif columns == 3:
        widths = [1.7, 2.1, 2.7]
    elif columns == 4:
        widths = [1.45, 1.7, 1.65, 1.7]
    for row_index, row_values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(_plain(value))
            _set_run_font(
                run,
                size=8.8,
                color=INK if row_index == 0 else "000000",
                bold=row_index == 0,
            )
            if row_index == 0:
                _set_cell_shading(cells[index], LIGHT_BLUE)
            elif row_index % 2 == 0:
                _set_cell_shading(cells[index], LIGHT_GRAY)
    _set_table_geometry(table, widths)
    _set_borders(table)
    _keep_table_rows_intact(table)


def _render_markdown(doc: Document, markdown: str):
    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        if not line.strip():
            index += 1
            continue
        if line.startswith("# "):
            index += 1
            while index < len(lines) and not lines[index].startswith("## "):
                index += 1
            continue
        if line.startswith("## "):
            doc.add_heading(_plain(line[3:]), level=1)
        elif line.startswith("### "):
            doc.add_heading(_plain(line[4:]), level=2)
        elif line.startswith("```"):
            index += 1
            code_lines = []
            while index < len(lines) and not lines[index].startswith("```"):
                code_lines.append(lines[index])
                index += 1
            paragraph = doc.add_paragraph(style="Code")
            for code_index, code_line in enumerate(code_lines):
                if code_index:
                    paragraph.add_run("\n")
                paragraph.add_run(code_line)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F4F6F9")
            paragraph._p.get_or_add_pPr().append(shading)
        elif line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            _add_table(doc, table_lines)
            continue
        elif line.startswith("- [ ] "):
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.add_run("☐ " + _plain(line[6:]))
        elif line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.add_run(_plain(line[2:]))
        elif len(line) > 3 and line[0].isdigit() and ". " in line[:4]:
            paragraph = doc.add_paragraph(style="List Number")
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.add_run(_plain(line.split(". ", 1)[1]))
        else:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.add_run(_plain(line.replace("  ", " ")))
        index += 1


def _plain(value: str) -> str:
    """Remove lightweight Markdown markers from the partner-facing document."""
    return value.replace("**", "").replace("`", "")


def main():
    document = Document()
    _configure_styles(document)
    _configure_section(document)
    _add_masthead(document)
    _render_markdown(document, SOURCE.read_text(encoding="utf-8"))
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
